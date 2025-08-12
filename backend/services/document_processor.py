import PyPDF2
import docx
import pandas as pd
from typing import Dict, List, Optional
import re
import logging
from pathlib import Path
import json
import asyncio
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, llm_analyzer):
        self.llm_analyzer = llm_analyzer
        self.supported_formats = ['.pdf', '.docx', '.txt', '.csv', '.xlsx', '.json']
        self.processed_documents = {}
        
    async def process_document(self, file_path: str, file_type: Optional[str] = None) -> Dict:
        """Process uploaded document and extract trading insights"""
        
        try:
            path = Path(file_path)
            
            if not path.exists():
                return {"error": f"File not found: {file_path}"}
            
            # Determine file type
            if not file_type:
                file_type = path.suffix.lower()
            
            if file_type not in self.supported_formats:
                return {"error": f"Unsupported file type: {file_type}"}
            
            # Extract text based on file type
            text_content = await self._extract_text(file_path, file_type)
            
            if not text_content:
                return {"error": "No text content extracted from document"}
            
            # Extract structured information
            extracted_data = self._extract_key_information(text_content)
            
            # Analyze with LLM
            llm_analysis = self.llm_analyzer.analyze_uploaded_report(text_content)
            
            # Store processed document
            doc_id = self._generate_doc_id(file_path)
            self.processed_documents[doc_id] = {
                'file_path': file_path,
                'processed_at': datetime.now().isoformat(),
                'insights': llm_analysis
            }
            
            return {
                'status': 'success',
                'doc_id': doc_id,
                'file_path': file_path,
                'file_type': file_type,
                'text_preview': text_content[:500],
                'extracted_data': extracted_data,
                'llm_analysis': llm_analysis,
                'processed_at': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {"error": str(e)}
    
    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from document based on file type"""
        
        try:
            if file_type == '.pdf':
                return await asyncio.to_thread(self._extract_pdf_text, file_path)
            elif file_type == '.docx':
                return await asyncio.to_thread(self._extract_docx_text, file_path)
            elif file_type == '.txt':
                return await asyncio.to_thread(self._extract_txt_text, file_path)
            elif file_type in ['.csv', '.xlsx']:
                return await asyncio.to_thread(self._extract_spreadsheet_data, file_path)
            elif file_type == '.json':
                return await asyncio.to_thread(self._extract_json_data, file_path)
            else:
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                        
            return '\n'.join(text)
            
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text.append(' | '.join(row_text))
            
            return '\n'.join(text)
            
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading TXT {file_path}: {e}")
                return ""
    
    def _extract_spreadsheet_data(self, file_path: str) -> str:
        """Extract data from CSV or Excel file"""
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Convert dataframe to readable text format
            text_parts = []
            
            # Add column descriptions
            text_parts.append(f"Columns: {', '.join(df.columns)}")
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_parts.append("\nNumeric Summary:")
                for col in numeric_cols:
                    text_parts.append(f"{col}: min={df[col].min():.2f}, max={df[col].max():.2f}, mean={df[col].mean():.2f}")
            
            # Add first few rows as sample
            text_parts.append("\nSample Data:")
            text_parts.append(df.head(10).to_string())
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error reading spreadsheet {file_path}: {e}")
            return ""
    
    def _extract_json_data(self, file_path: str) -> str:
        """Extract data from JSON file"""
        try:
            with open(file_path, 'r') as file:
                data = json.load(file)
            
            # Convert JSON to readable format
            return json.dumps(data, indent=2)
            
        except Exception as e:
            logger.error(f"Error reading JSON {file_path}: {e}")
            return ""
    
    def _extract_key_information(self, text: str) -> Dict:
        """Extract structured information from text"""
        
        extracted = {
            'prices': [],
            'percentages': [],
            'dates': [],
            'tickers': [],
            'keywords': {},
            'patterns': {}
        }
        
        # Define patterns
        patterns = {
            'prices': r'\$?(\d+(?:,\d+)*(?:\.\d+)?)\s*(?:USD|KRW|원)?',
            'percentages': r'([+-]?\d+(?:\.\d+)?)\s*%',
            'dates': r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})',
            'tickers': r'\b([A-Z]{3,5}(?:-[A-Z]{3,5})?)\b',
            'time_periods': r'(\d+)\s*(day|week|month|year|hour|minute)s?',
            'support_resistance': r'(?:support|resistance|target)\s*(?:at|level|price)?\s*\$?(\d+(?:,\d+)*(?:\.\d+)?)'
        }
        
        # Extract prices
        price_matches = re.findall(patterns['prices'], text)
        extracted['prices'] = [self._clean_number(p) for p in price_matches[:20]]
        
        # Extract percentages
        percent_matches = re.findall(patterns['percentages'], text)
        extracted['percentages'] = [float(p) for p in percent_matches[:20]]
        
        # Extract dates
        extracted['dates'] = re.findall(patterns['dates'], text)[:10]
        
        # Extract potential tickers
        ticker_matches = re.findall(patterns['tickers'], text)
        # Filter common words that match pattern but aren't tickers
        common_words = {'THE', 'AND', 'FOR', 'WITH', 'FROM', 'THIS', 'THAT'}
        extracted['tickers'] = list(set(t for t in ticker_matches if t not in common_words))[:20]
        
        # Extract time periods
        time_matches = re.findall(patterns['time_periods'], text)
        extracted['time_periods'] = time_matches[:10]
        
        # Extract support/resistance levels
        sr_matches = re.findall(patterns['support_resistance'], text, re.IGNORECASE)
        extracted['support_resistance'] = [self._clean_number(sr) for sr in sr_matches[:10]]
        
        # Count important keywords
        keywords = [
            'buy', 'sell', 'hold', 'long', 'short',
            'bullish', 'bearish', 'neutral',
            'support', 'resistance', 'breakout', 'breakdown',
            'target', 'stop', 'loss', 'profit',
            'trend', 'reversal', 'continuation',
            'overbought', 'oversold',
            'accumulation', 'distribution'
        ]
        
        text_lower = text.lower()
        for keyword in keywords:
            count = len(re.findall(rf'\b{keyword}\b', text_lower))
            if count > 0:
                extracted['keywords'][keyword] = count
        
        # Identify potential patterns
        pattern_indicators = {
            'head_and_shoulders': r'head\s*and\s*shoulders',
            'double_top': r'double\s*top',
            'double_bottom': r'double\s*bottom',
            'triangle': r'(?:ascending|descending|symmetrical)\s*triangle',
            'flag': r'(?:bull|bear)\s*flag',
            'wedge': r'(?:rising|falling)\s*wedge',
            'cup_and_handle': r'cup\s*and\s*handle'
        }
        
        for pattern_name, pattern_regex in pattern_indicators.items():
            if re.search(pattern_regex, text_lower):
                extracted['patterns'][pattern_name] = True
        
        return extracted
    
    def _clean_number(self, number_str: str) -> float:
        """Clean and convert number string to float"""
        try:
            # Remove commas and currency symbols
            cleaned = re.sub(r'[,$]', '', number_str)
            return float(cleaned)
        except:
            return 0.0
    
    def _generate_doc_id(self, file_path: str) -> str:
        """Generate unique document ID"""
        import hashlib
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        hash_obj = hashlib.md5(file_path.encode())
        return f"DOC_{timestamp}_{hash_obj.hexdigest()[:8]}"
    
    def get_processed_documents(self) -> List[Dict]:
        """Get list of processed documents"""
        return list(self.processed_documents.values())
    
    def get_document_insights(self, doc_id: str) -> Optional[Dict]:
        """Get insights for a specific document"""
        return self.processed_documents.get(doc_id)
    
    async def batch_process_documents(self, file_paths: List[str]) -> List[Dict]:
        """Process multiple documents in batch"""
        tasks = [self.process_document(fp) for fp in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        processed = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed.append({
                    "error": str(result),
                    "file_path": file_paths[i]
                })
            else:
                processed.append(result)
        
        return processed
    
    def combine_document_insights(self, doc_ids: List[str]) -> Dict:
        """Combine insights from multiple documents"""
        
        combined_insights = {
            'price_targets': [],
            'trends': [],
            'risks': [],
            'opportunities': [],
            'recommendations': [],
            'consensus': None
        }
        
        for doc_id in doc_ids:
            doc = self.processed_documents.get(doc_id)
            if doc and 'insights' in doc:
                insights = doc['insights']
                
                if 'price_targets' in insights:
                    combined_insights['price_targets'].extend(insights['price_targets'])
                if 'trends' in insights:
                    combined_insights['trends'].extend(insights['trends'])
                if 'risks' in insights:
                    combined_insights['risks'].extend(insights['risks'])
                if 'recommendations' in insights:
                    combined_insights['recommendations'].append(insights['recommendations'])
        
        # Determine consensus
        if combined_insights['recommendations']:
            rec_counts = {}
            for rec in combined_insights['recommendations']:
                if isinstance(rec, str):
                    rec_counts[rec] = rec_counts.get(rec, 0) + 1
            
            if rec_counts:
                combined_insights['consensus'] = max(rec_counts, key=rec_counts.get)
        
        return combined_insights